#!/usr/bin/env python3
"""Build a formatted .docx project report for FleetOps."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG = "/home/user/zoho-linux-fleet-manager/fleetops.report/img"
OUT = "/home/user/zoho-linux-fleet-manager/fleetops.report/FleetOps-Project-Report.docx"

DARK = RGBColor(0x1e, 0x29, 0x3b)
ACCENT = RGBColor(0x2563, 0x000, 0x000) if False else RGBColor(0x1f, 0x6f, 0xeb)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _no_borders_except_bottom(t, line_hex="dbe3ec"):
    """Clean borderless table with thin horizontal row separators only."""
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    inH = OxmlElement("w:insideH")
    inH.set(qn("w:val"), "single")
    inH.set(qn("w:sz"), "4")
    inH.set(qn("w:space"), "0")
    inH.set(qn("w:color"), line_hex)
    borders.append(inH)
    tblPr.append(borders)


def add_table(headers, rows, widths=None, accent="2F5FBF", header_text="FFFFFF",
              zebra="F4F7FB"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders_except_bottom(t)
    # header
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        _set_cell_margins(hdr[i], top=90, bottom=90, left=150, right=150)
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(header_text)
        set_cell_bg(hdr[i], accent)
    # body
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            _set_cell_margins(cells[i], top=80, bottom=80, left=150, right=150)
            if ri % 2 == 1:
                set_cell_bg(cells[i], zebra)
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x2a, 0x2f, 0x37)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(accent)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, size=11, italic=False, color=None, align=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def img(path, width=6.3):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def caption(text):
    p = para(text, size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p

def _set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def _accent_bar(cell, hexcolor):
    """Left vertical color bar via cell left border."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "36")  # thick
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hexcolor)
    borders.append(left)
    tcPr.append(borders)

def card(icon, title, accent_hex, fill_hex, body):
    """A colored callout card: accent left bar, tinted fill, bold icon+title, body text."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_bg(cell, fill_hex)
    _accent_bar(cell, accent_hex)
    _set_cell_margins(cell)
    # title line
    cell.text = ""
    pt = cell.paragraphs[0]
    ri = pt.add_run(f"{icon}  ")
    ri.bold = True
    ri.font.size = Pt(13)
    ri.font.color.rgb = RGBColor.from_string(accent_hex)
    rt = pt.add_run(title)
    rt.bold = True
    rt.font.size = Pt(12)
    rt.font.color.rgb = RGBColor.from_string(accent_hex)
    # body line
    pb = cell.add_paragraph()
    rb = pb.add_run(body)
    rb.font.size = Pt(10)
    rb.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph()  # spacing
    return t

def links_box(items):
    """Centered box on cover listing project access links."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(5.2)
    set_cell_bg(cell, "F4F6FB")
    _accent_bar(cell, "1F6FEB")
    _set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    cell.text = ""
    head = cell.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = head.add_run("PROJECT LINKS")
    hr.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = RGBColor.from_string("1F6FEB")
    for label, val in items:
        p = cell.add_paragraph()
        lr = p.add_run(f"{label}:  ")
        lr.bold = True
        lr.font.size = Pt(10.5)
        vr = p.add_run(val)
        vr.font.size = Pt(10.5)
        vr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    doc.add_paragraph()

# ---------- COVER ----------
for _ in range(2):
    doc.add_paragraph()
para("FleetOps", size=40, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK)
para("A Web-Based Linux Fleet Management System\nwith Prometheus & Grafana Monitoring",
     size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_paragraph()
para("PROJECT REPORT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT)
for _ in range(2):
    doc.add_paragraph()

fields = [
    ("Submitted by", "Divyani Khandait"),
    ("Roll No. / ID", "____________________________"),
    ("Department", "____________________________"),
    ("College / Institution", "____________________________"),
    ("Guide / Mentor", "____________________________"),
    ("Date", "____________________________"),
]
for label, val in fields:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(f"{label}:  ")
    rb.bold = True
    rb.font.size = Pt(12)
    rv = p.add_run(val)
    rv.font.size = Pt(12)

doc.add_paragraph()
links_box([
    ("GitHub Repository", "https://github.com/Divyaaaani/Linux_fleet-3621"),
    ("Live Demo", "____________________  (run locally — see Section 11)"),
])
para("Zoho SETU — Project 3: Distributed Linux Fleet Management",
     size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_page_break()

# ---------- 1 EXEC SUMMARY ----------
h("1. Executive Summary", 1)
para("FleetOps is a web-based control plane that onboards, monitors, and remotely manages a fleet "
     "of Linux machines from a single dashboard, and exposes its telemetry through industry-standard "
     "Prometheus and Grafana monitoring.")
para("Managing many Linux servers by hand does not scale: there is no central view of health, SSH "
     "access is painful to configure, machines behind firewalls/NAT are unreachable for pull-based "
     "tools, failures go unnoticed, and ad-hoc remote shell access is dangerous. FleetOps solves all "
     "five problems with one coherent design — a lightweight push-based bash agent on each node and a "
     "central control plane that stores metrics, raises alerts, dispatches safe (allow-listed) "
     "commands, and exports everything to Prometheus for Grafana to visualise.")
para("The system fully satisfies the project brief: it is web-based, it onboards and manages multiple "
     "Linux nodes remotely, and it collects, stores, and visualises system metrics using Prometheus "
     "and Grafana.")

# ---------- 2 OBJECTIVES ----------
h("2. Objectives", 1)
add_table(["#", "Objective", "Delivered"], [
    ["1", "Web-based interface to manage a Linux fleet", "React dashboard in the browser"],
    ["2", "Onboard multiple Linux nodes remotely", "Token-based, single-command agent install"],
    ["3", "Manage nodes remotely", "Allow-listed remote command dispatch"],
    ["4", "Collect system metrics", "Bash agent reads /proc every 15s"],
    ["5", "Store metrics", "SQLite (Turso) time-series tables"],
    ["6", "Visualise metrics", "Custom charts + Grafana dashboards"],
    ["7", "Use industry-standard monitoring", "Prometheus exporter + Grafana"],
], widths=[0.4, 3.2, 3.0])

# ---------- 3 DESIGN PRIORITIES ----------
h("3. Design Priorities", 1)
para("This project deliberately prioritises correctness, system design, observability, and an "
     "understanding of Linux internals over UI aesthetics. Each priority is backed by a concrete "
     "design choice:")
doc.add_paragraph()

card("✔", "Correctness", "16A34A", "EAF7EE",
     "One-time tokens consumed atomically at registration; idempotent metric ingest; a state machine "
     "(pending → online → degraded/offline) driven by heartbeat freshness; commands tracked through "
     "queued → running → success/failure with results written back. The control plane is the single "
     "source of truth.")
card("◇", "System Design", "2563EB", "E8F0FE",
     "Clear separation of agent / control plane / presentation tiers. A push-based agent that works "
     "behind NAT and firewalls, an allow-listed (never arbitrary) command channel, and a normalised "
     "relational schema. Trade-offs are explicit and defensible.")
card("◎", "Observability", "D97706", "FEF3E2",
     "A native Prometheus exporter in standard exposition format, scraped into Prometheus and "
     "visualised in Grafana — plus an automatic alert engine for threshold breaches and missed "
     "heartbeats. Every node is measurable end-to-end.")
card("⌘", "Linux Internals", "7C3AED", "F1EBFC",
     "The agent reads directly from the kernel: CPU from /proc/stat, memory from /proc/meminfo, load "
     "from /proc/loadavg, uptime from /proc/uptime, disk via df, and process count from /proc. Pure "
     "bash + coreutils, with no runtime dependencies.")

doc.add_paragraph()
note = doc.add_paragraph()
ni = note.add_run("Note:  ")
ni.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
nr = note.add_run("UI aesthetics were intentionally kept functional rather than decorative — the dark "
                  "operations-centre theme exists only to make dense telemetry readable, not as a design "
                  "goal in itself.")
nr.italic = True
nr.font.size = Pt(10)
nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ---------- 4 ARCHITECTURE ----------
h("4. System Architecture", 1)
para("The system has three logical tiers: the fleet (agents), the control plane (API, storage, "
     "alerting, exporter), and the presentation layer (web dashboard + Grafana).")
img(f"{IMG}/diagram-architecture.png", 6.4)
caption("Figure 1 — FleetOps system architecture")

h("4.1 Components", 2)
for b in [
    "Agent (agent.sh) — a dependency-free bash script on every node; reads CPU, memory, disk, load, "
    "uptime and process count from /proc and reports over HTTP.",
    "Control plane — a Hono REST API on the Bun runtime, backed by Drizzle ORM over SQLite/Turso. "
    "Registers nodes, ingests telemetry, evaluates alerts, queues commands, exposes a Prometheus exporter.",
    "Web dashboard — a React + Tailwind single-page app polling every 5 seconds for a live feel.",
    "Prometheus — scrapes the control plane's exporter and stores metrics as a time series.",
    "Grafana — queries Prometheus and renders dashboards; also embedded inside the FleetOps UI.",
]:
    doc.add_paragraph(b, style="List Bullet")

h("4.2 The key design decision: push, not pull", 2)
para("The agent pushes metrics to the control plane and pulls queued commands back in the same HTTP "
     "request. A node never needs an open inbound port — only the control plane must be reachable. "
     "As a result the system works for machines behind NAT or firewalls, exactly where traditional "
     "pull-based monitoring fails. Because the agent is push-based, Prometheus scrapes the control "
     "plane's aggregated exporter, not each node — keeping configuration to a single target.")

h("4.3 Safety by design", 2)
para("Remote control is restricted to an explicit allow-list of actions — refresh metrics, health "
     "check, restart agent, restart a named service. The system never executes arbitrary shell sent "
     "over the wire, so remote management is safe by construction.")

# ---------- 4 LIFECYCLE ----------
h("5. Node Lifecycle & Data Flow", 1)
img(f"{IMG}/diagram-flow.png", 4.7)
caption("Figure 2 — Node lifecycle: onboard → report → control")
steps = [
    "An administrator generates a one-time token in the dashboard.",
    "The operator runs agent.sh on the target node with that token (single curl | bash command).",
    "The agent calls POST /api/agent/register with host facts; the control plane validates the token and creates the node record.",
    "The agent loops, calling POST /api/agent/:id/report every 15 seconds with fresh metrics.",
    "The alert + health engine stores each sample and evaluates thresholds and heartbeat freshness.",
    "The report response carries queued commands; the agent executes the allow-listed action and reports the result.",
    "The dashboard and Grafana visualise live state, time-series charts, and alerts.",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

# ---------- 5 DATA MODEL ----------
h("6. Data Model", 1)
add_table(["Table", "Purpose"], [
    ["nodes", "Registered machines, host facts, current status, last heartbeat"],
    ["onboarding_tokens", "One-time tokens consumed at registration"],
    ["metrics", "Time-series telemetry per node"],
    ["commands", "Queued / running / completed remote actions"],
    ["alerts", "Threshold + heartbeat alerts, with resolution timestamps"],
], widths=[1.8, 4.5])
para("A node moves through pending → online, and may become degraded (a threshold breach) or offline "
     "(missed heartbeats), automatically returning to online when healthy again.")

# ---------- 6 API ----------
h("7. REST API", 1)
add_table(["Method", "Route", "Purpose"], [
    ["POST", "/api/tokens", "Generate onboarding token"],
    ["POST", "/api/agent/register", "Agent registers with token"],
    ["POST", "/api/agent/:id/report", "Heartbeat + metrics; returns queued commands"],
    ["POST", "/api/agent/:id/command-result", "Agent reports command result"],
    ["GET", "/api/nodes", "Fleet list with latest metrics"],
    ["GET", "/api/nodes/:id/metrics", "Time-series (last hour)"],
    ["POST", "/api/nodes/:id/commands", "Queue a remote command"],
    ["GET", "/api/alerts", "Active + resolved alerts"],
    ["GET", "/api/summary", "Fleet health counts"],
    ["GET", "/api/metrics/prometheus", "Prometheus exposition format for scraping"],
], widths=[0.8, 2.6, 2.9])

# ---------- 7 PROM/GRAFANA ----------
h("8. Prometheus & Grafana Integration", 1)
h("8.1 Exporter", 2)
para("The control plane exposes GET /api/metrics/prometheus in the standard Prometheus text exposition "
     "format. Each metric is labelled with node, ip, env, and os for filtering and aggregation in Grafana.")
code = ('# HELP fleet_cpu_percent CPU utilization percent\n'
        '# TYPE fleet_cpu_percent gauge\n'
        'fleet_node_up{node="web-prod-01",ip="10.0.109.188",env="prod",os="Rocky Linux"} 1\n'
        'fleet_cpu_percent{node="web-prod-01",...} 43.7\n'
        'fleet_memory_percent{node="web-prod-01",...} 41.6\n'
        'fleet_disk_percent{node="web-prod-01",...} 72\n'
        'fleet_load1{node="web-prod-01",...} 0.16')
cp = doc.add_paragraph()
cr = cp.add_run(code)
cr.font.name = "Consolas"
cr.font.size = Pt(8.5)
para("Exported series include fleet_node_up, fleet_cpu_percent, fleet_memory_percent, "
     "fleet_disk_percent, fleet_load1, fleet_uptime_seconds, fleet_processes, fleet_alerts_open, "
     "and fleet_nodes_total.")

h("8.2 One-command stack", 2)
para("A docker-compose.yml brings up the full monitoring stack — the app, Prometheus, and Grafana — "
     "together. Prometheus scrapes the exporter every 15s, and Grafana is provisioned automatically "
     "with the Prometheus datasource and a ready-made fleet dashboard. Grafana is also embedded "
     "directly inside the FleetOps UI.")
cp = doc.add_paragraph()
cr = cp.add_run("docker compose up -d\n# App        → http://localhost:4200\n"
                "# Prometheus → http://localhost:9090\n# Grafana    → http://localhost:3000")
cr.font.name = "Consolas"
cr.font.size = Pt(9)

# ---------- 8 SCREENSHOTS ----------
h("9. Implementation Screenshots", 1)
h("9.1 Fleet Overview", 2)
para("Real-time health cards, a live node list with per-node CPU/memory bars, and an active-alerts panel.")
img(f"{IMG}/01-overview.png", 6.4)
caption("Figure 3 — Fleet overview dashboard")
doc.add_page_break()

h("9.2 Node Detail", 2)
para("Full system facts, live resource bars, time-series charts, and a Remote Actions panel with command history.")
img(f"{IMG}/03-node-detail.png", 5.6)
caption("Figure 4 — Node detail with live charts and remote actions")
doc.add_page_break()

h("9.3 Fleet — Nodes View", 2)
para("Searchable, filterable table of all managed nodes with status, environment, and latest metrics.")
img(f"{IMG}/02-nodes.png", 6.4)
caption("Figure 5 — Nodes table")

h("9.4 Onboarding", 2)
para("Generate a one-time token and copy the single-command agent installer.")
img(f"{IMG}/05-onboarding.png", 6.4)
caption("Figure 6 — Node onboarding")
doc.add_page_break()

# ---------- 9 STACK ----------
h("10. Technology Stack", 1)
add_table(["Layer", "Technology"], [
    ["Agent", "Bash + coreutils (no dependencies)"],
    ["Runtime", "Bun"],
    ["API", "Hono (REST)"],
    ["ORM / DB", "Drizzle ORM over SQLite (Turso)"],
    ["Frontend", "React 19 + Wouter + Tailwind CSS"],
    ["Monitoring", "Prometheus + Grafana"],
    ["Orchestration", "Docker Compose"],
], widths=[2.0, 4.3])

# ---------- 10 RUN ----------
h("11. How to Run", 1)
runcode = ('# 1. Install dependencies\nbun install\n\n'
           '# 2. Configure database (.env: DATABASE_URL, DATABASE_AUTH_TOKEN)\n'
           'cd packages/web && bun run db:push && cd ../..\n\n'
           '# 3. Start the control plane + dashboard\n'
           'bun run dev --port 4200          # http://localhost:4200\n\n'
           '# 4. Onboard a real node\n'
           'curl -fsSL http://<server>:4200/agent.sh | sudo bash -s -- \\\n'
           '  --server http://<server>:4200 --token flt_xxxxxxxx\n\n'
           '# 4b. ...or simulate a whole fleet for testing\n'
           'bash simulate-fleet.sh 7\n\n'
           '# 5. Bring up Prometheus + Grafana\n'
           'docker compose up -d')
cp = doc.add_paragraph()
cr = cp.add_run(runcode)
cr.font.name = "Consolas"
cr.font.size = Pt(9)

# ---------- 11 RESULTS ----------
h("12. Results & Conclusion", 1)
para("FleetOps delivers a complete, working Linux fleet management system that meets every requirement "
     "of the brief:")
for b in [
    "Web-based management of a multi-node Linux fleet from one dashboard.",
    "Remote onboarding in a single command, working even behind NAT/firewalls.",
    "Remote management through a safe, allow-listed command channel.",
    "Metric collection from each node via a lightweight bash agent.",
    "Storage of time-series telemetry in SQLite/Turso.",
    "Visualisation through both a custom real-time dashboard and Prometheus + Grafana.",
]:
    doc.add_paragraph(b, style="List Bullet")
para("The defining engineering choice — a push-based agent — is what makes the system practical for "
     "real infrastructure, where servers commonly sit behind firewalls. Combined with allow-listed "
     "remote control and standard Prometheus/Grafana observability, FleetOps is a small but "
     "production-shaped take on distributed Linux fleet management.")

h("Future work", 2)
for b in [
    "Role-based access control and audit logging for commands.",
    "Alerting integrations (email / Slack / PagerDuty) via Prometheus Alertmanager.",
    "Agent auto-update and rolling command rollout across the fleet.",
    "Historical metric retention beyond the live window.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.save(OUT)
print("saved", OUT)
